"""
Page 6: Report Generator
Generate session summaries grouped by regulation
"""
import streamlit as st
from supabase_client import SupabaseClient
from gemini_client import GeminiClient
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from datetime import datetime

st.set_page_config(page_title="Report Generator", page_icon="📊", layout="wide")

from auth_utils import require_auth
require_auth()

st.title("📊 Report Generator")
st.markdown("Generate AI-powered session reports grouped by regulation")

# ============================================================================
# SESSION SELECTOR
# ============================================================================

st.markdown("### 1️⃣ Select Session")

col1, col2 = st.columns(2)

with col1:
    try:
        groups = SupabaseClient.get_all_groups()
        if not groups:
            st.error("No groups found")
            st.stop()
        
        selected_group = st.selectbox(
            "Group",
            options=[g['id'] for g in groups],
            format_func=lambda x: f"{x} - {next(g['full_name'] for g in groups if g['id'] == x)}"
        )
    except Exception as e:
        st.error(f"Error: {e}")
        st.stop()

with col2:
    try:
        sessions = SupabaseClient.get_sessions_by_group(selected_group)
        if not sessions:
            st.error("No sessions found")
            st.stop()
        
        selected_session = st.selectbox(
            "Session",
            options=sessions,
            format_func=lambda x: f"Session {x['code']} ({x['year']})"
        )
    except Exception as e:
        st.error(f"Error: {e}")
        st.stop()

# ============================================================================
# GENERATE REPORT
# ============================================================================

st.markdown("---")
st.markdown("### 2️⃣ Generate Report")

if st.button("🤖 Generate Summary", type="primary", use_container_width=True):
    with st.spinner("Generating report..."):
        try:
            # Fetch all documents from session
            documents = SupabaseClient.get_documents_by_session(selected_session['id'])
            
            if not documents:
                st.warning("No documents found in this session")
                st.stop()
            
            # Group by regulation
            by_regulation = {}
            for doc in documents:
                reg_id = doc.get('regulation_ref_id', 'General Topics')
                if reg_id not in by_regulation:
                    by_regulation[reg_id] = []
                by_regulation[reg_id].append(doc)
            
            # Generate summary
            st.markdown("### 📄 Generated Summary")
            
            full_summary = f"# Session Report: {selected_group} Session {selected_session['code']} ({selected_session['year']})\n\n"
            
            if selected_session.get('dates'):
                full_summary += f"**Dates:** {selected_session['dates']}\n\n"
            
            full_summary += "---\n\n"
            
            # Generate for each regulation
            for reg_id, docs in by_regulation.items():
                st.markdown(f"#### Processing: {reg_id}")
                
                # Build document list
                doc_list = "\n".join([
                    f"- **{doc['symbol']}**: {doc['title']} (by {doc['author']})"
                    for doc in docs
                ])
                
                # Use Gemini to summarize
                prompt = f"""
Summarize the key discussions and decisions for regulation {reg_id} based on these documents from {selected_group} Session {selected_session['code']}:

{doc_list}

Provide a structured summary covering:
1. Main topics discussed
2. Key proposals and authors
3. Decisions or outcomes (if mentioned)
4. Next steps or pending items

Be concise but comprehensive. Use bullet points for clarity.
"""
                
                try:
                    model = GeminiClient.generate_embedding.__self__  # Access to genai
                    import google.generativeai as genai
                    model = genai.GenerativeModel('gemini-1.5-pro')
                    response = model.generate_content(prompt)
                    summary_text = response.text
                except Exception as e:
                    summary_text = f"Error generating summary: {e}"
                
                # Add to full summary
                section = f"## {reg_id}\n\n{summary_text}\n\n"
                full_summary += section
                
                # Display in UI
                with st.expander(f"📋 {reg_id}", expanded=True):
                    st.markdown(summary_text)
            
            # Store in session state for download
            st.session_state.generated_summary = full_summary
            st.session_state.session_info = {
                'group': selected_group,
                'code': selected_session['code'],
                'year': selected_session['year']
            }
            
            st.success("✅ Report generated successfully!")
        
        except Exception as e:
            st.error(f"Error generating report: {e}")
            st.exception(e)

# ============================================================================
# DOWNLOAD AS WORD
# ============================================================================

if 'generated_summary' in st.session_state:
    st.markdown("---")
    st.markdown("### 3️⃣ Download Report")
    
    col1, col2 = st.columns(2)
    
    with col1:
        # Download as markdown
        st.download_button(
            label="📥 Download as Markdown",
            data=st.session_state.generated_summary,
            file_name=f"{st.session_state.session_info['group']}_Session_{st.session_state.session_info['code']}_Report.md",
            mime="text/markdown",
            use_container_width=True
        )
    
    with col2:
        # Generate Word document
        if st.button("📄 Generate Word Document", use_container_width=True):
            try:
                # Create Word document
                doc = Document()
                
                # Title
                title = doc.add_heading(
                    f"{st.session_state.session_info['group']} Session {st.session_state.session_info['code']} Report",
                    level=0
                )
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Subtitle
                subtitle = doc.add_paragraph(f"Year: {st.session_state.session_info['year']}")
                subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                doc.add_paragraph()
                
                # Add summary content (parse markdown to add headings)
                lines = st.session_state.generated_summary.split('\n')
                for line in lines:
                    if line.startswith('## '):
                        doc.add_heading(line[3:], level=1)
                    elif line.startswith('# '):
                        doc.add_heading(line[2:], level=0)
                    elif line.strip():
                        doc.add_paragraph(line)
                
                # Save to bytes
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                # Download button
                st.download_button(
                    label="📥 Download Word Document",
                    data=buffer,
                    file_name=f"{st.session_state.session_info['group']}_Session_{st.session_state.session_info['code']}_Report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                
                st.success("✅ Word document ready for download!")
            
            except Exception as e:
                st.error(f"Error creating Word document: {e}")
